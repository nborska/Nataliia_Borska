from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)

def add_warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

def add_good(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x37, 0x86, 0x30)
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
    doc.add_paragraph()

# ===== ДОКУМЕНТ =====

add_title(doc, 'АУДИТ EMAIL-СІКВЕНСІВ MSN')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AIR Media-Tech  |  MSN Distribution  |  Травень 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
doc.add_paragraph()

# 1. ЗАГАЛЬНИЙ ВИСНОВОК
add_h1(doc, '1. ЗАГАЛЬНИЙ ВИСНОВОК')
add_body(doc, 'З 13 сіквенсів — 2-3 сильні, решта потребує правок. Найбільша проблема: непослідовні факти, штучна терміновість і надто багато CTA-варіантів в одному листі.')
doc.add_paragraph()

# 2. КРИТИЧНІ ПОМИЛКИ
add_h1(doc, '2. КРИТИЧНІ ПОМИЛКИ — ВИПРАВИТИ ПЕРЕД ВІДПРАВКОЮ')
doc.add_paragraph()

add_h2(doc, 'Помилка 1 — Різні цифри аудиторії')
add_body(doc, 'В різних сіквенсах вказано "2B users" (Launch 1, 3, 4, 5, 6, 7, 8, 9, 12, 13) і "700M monthly users" (Launch 10).')
add_body(doc, 'Правильна цифра за гайдлайном: 700M monthly users (MSN сайт + app). Цифра 2B — це загальна аудиторія всіх Microsoft products, що є перебільшенням для MSN.')
add_warn(doc, 'Якщо одна людина отримає два листи з різними цифрами — довіра до відправника нульова.')
add_good(doc, 'Рішення: замінити скрізь на "700M+ monthly users across Bing, Edge, Windows feeds, and the MSN app"')
doc.add_paragraph()

add_h2(doc, 'Помилка 2 — Launch 5: штучна терміновість "50 запрошень"')
add_body(doc, 'Текст листа: "Ми отримали 50 додаткових запрошень, перебрали багато каналів... від Microsoft ваучери тільки на 50 запрошень"')
add_warn(doc, 'Якщо це перебільшення — це репутаційний ризик. Такий підхід може спрацювати один раз, але вбиває довіру при масштабуванні.')
add_good(doc, 'Рішення: видалити Launch 5 або перефразувати чесно — MSN доступний тільки через офіційних партнерів. Це реальне обмеження, не маркетинговий прийом.')
doc.add_paragraph()

add_h2(doc, 'Помилка 3 — 5-6 варіантів CTA в одному листі')
add_body(doc, 'Майже кожен лист закінчується кількома варіантами: А) спробуємо? В) що скажеш? С) коли зідзвонимось? Д) зацікавлен? Е) домовились?')
add_warn(doc, 'Це чернетки з варіантами для вибору, а не фінальні листи. Більше варіантів = менше конверсія (decision paralysis).')
add_good(doc, 'Рішення: перед запуском залишити рівно ОДИН CTA на кожен лист.')
doc.add_paragraph()

add_h2(doc, 'Помилка 4 — Launch 7: два продукти в одному листі')
add_body(doc, 'Текст: "1 — Можу допомогти потрапити на MSN... 2 — Можу підключити наш Safety бандл..."')
add_warn(doc, 'В холодному аутрічі два продукти в першому листі заплутують читача. Людина не розуміє що саме їй пропонують.')
add_good(doc, 'Рішення: один лист — один оффер. MSN окремо, Safety Bundle окремо.')
doc.add_paragraph()

add_h2(doc, 'Помилка 5 — Launch 3: хайп і порушення гайдлайну')
add_body(doc, 'Текст: "I\'m excited to offer you an exclusive invitation... revolutionizing the world of content creation!"')
add_warn(doc, 'Порушує гайдлайн: немає конкретних цифр, є слово "revolutionary", є зайвий ентузіазм без доказів.')
add_good(doc, 'Рішення: вивести Launch 3 з ротації або повністю переписати.')
doc.add_paragraph()

# 3. РЕЙТИНГ
add_h1(doc, '3. РЕЙТИНГ ВСІХ СІКВЕНСІВ')
doc.add_paragraph()

add_table(doc,
    ['Сіквенс', 'Оцінка', 'Коментар'],
    [
        ['Launch 10', '5/5 — Найкращий', '"Same videos. Second paycheck." — точно по гайдлайну, коротко, з доказами, 700M коректно'],
        ['Launch 1B', '4/5 — Добрий', 'Хороша структура і кейси. Треба виправити "2B" на "700M"'],
        ['Launch 4', '4/5 — Добрий', 'Loss frame — актуальний, Algorithm Apocalypse добре відчувається'],
        ['Launch 8', '4/5 — Добрий', 'Сегментація growing/declining — розумний підхід'],
        ['Launch 9', '3/5 — Середній', 'Гарна ідея з Brazilian Portuguese, але потребує доопрацювання'],
        ['Launch 13', '3/5 — Середній', '"Plan B" angle — сильний, але "10 slots this month" — штучна терміновість'],
        ['Launch 11', '3/5 — Середній', 'Аудит + MSN — цікавий підхід, але складно масштабувати'],
        ['Launch 12', '3/5 — Середній', 'Country personalization — гарна ідея, але сіквенс занадто короткий'],
        ['Launch 1A', '3/5 — Середній', 'Занадто довгий перший лист'],
        ['Launch 6', '2/5 — Слабкий', 'Нічого виразного, середній по всіх критеріях'],
        ['Launch 2', '2/5 — Слабкий', 'Занадто коротко і порожньо, немає доказів'],
        ['Launch 7', '2/5 — Слабкий', 'Мікс двох продуктів в одному листі — заплутує'],
        ['Launch 3', '1/5 — Вивести', 'Хайп, шаблонний спам, порушує гайдлайн'],
        ['Launch 5', 'Репутаційний ризик', 'Штучна терміновість — може зруйнувати довіру до AIR'],
    ]
)

# 4. A/B ТЕСТ
add_h1(doc, '4. ЩО ТЕСТУВАТИ ЗАРАЗ — A/B ТЕСТ (Exp. №2 З ПЛАНУ)')
doc.add_paragraph()
add_body(doc, 'З плану: Експеримент №2 — A/B тест Loss Frame vs Gain Frame (ICE Score 8.3). Найкращі кандидати:')
doc.add_paragraph()

add_h2(doc, 'Варіант A — Gain Frame (контроль): Launch 10')
add_body(doc, 'Subject: Same videos. Second paycheck for {{CHANNELNAME}}')
add_body(doc, 'Фокус: що ти ОТРИМАЄШ — +20-30% до доходу, пасивно, вже наявні відео.')
doc.add_paragraph()

add_h2(doc, 'Варіант B — Loss Frame (тест): Launch 4 адаптований')
add_body(doc, 'Subject: YouTube algorithms are constantly changing')
add_body(doc, 'Фокус: що ти ЗАРАЗ ВТРАЧАЄШ — твої відео не заробляють поки YouTube нестабільний.')
doc.add_paragraph()

add_body(doc, 'Метрика: reply rate і конверсія по кожній групі після мінімум 200 відправок.')
doc.add_paragraph()

# 5. НІШОВІ ВЕРСІЇ
add_h1(doc, '5. ПРІОРИТЕТ: ЗРОБИТИ 7 НІШОВИХ ВЕРСІЙ LAUNCH 10')
add_body(doc, 'Launch 10 — найсильніший шаблон. Потрібно 7 версій де змінюється тільки один кейс:')
doc.add_paragraph()

add_table(doc,
    ['Ніша', 'Кейс для підстановки'],
    [
        ['Food & Cooking', 'A cooking channel added approx. $1,300/month from its back catalog alone'],
        ['DIY & Crafts', 'A DIY channel pulled 17M views in 3 months from older tutorials'],
        ['Tech & Science', 'A tech channel went from 1.5M to 6.7M MSN views in 3 months'],
        ['Handmade & Jewelry', 'A jewelry channel got 1.3M MSN views in under 2 months, equal to lifetime YouTube views'],
        ['Health & Lifestyle', 'A lifehacks channel added +1.8M views in less than 2 months'],
        ['Entertainment', 'An action comedy channel added +677K monthly MSN views on top of YouTube'],
        ['Tech News & Finance', 'A tech news back catalog brought in approx. +$1,500 in a couple of months'],
    ]
)

# 6. СТРУКТУРА ІДЕАЛЬНОГО ЛИСТА
add_h1(doc, '6. СТРУКТУРА ІДЕАЛЬНОГО ПЕРШОГО ЛИСТА (на базі Launch 10)')
doc.add_paragraph()

add_h2(doc, 'Рядок 1 — Персоналізація і спостереження')
add_body(doc, 'Приклад: "I came across your channel and noticed you have strong evergreen [НІША] content."')

add_h2(doc, 'Рядок 2 — Хто ми і що робимо (одне речення)')
add_body(doc, 'Приклад: "We work with creators to distribute their existing videos across Microsoft surfaces — Bing, Edge, Windows feeds, and the MSN app — reaching over 700M monthly users."')

add_h2(doc, 'Рядок 3 — Що від тебе не потрібно')
add_body(doc, 'Приклад: "No need to film or upload anything new. We work with the videos you already have."')

add_h2(doc, 'Рядок 4 — Соціальний доказ з цифрами (нішовий кейс)')
add_body(doc, 'Приклад для Food: "A cooking channel recently generated an additional ~$1,300/month from its back catalog alone."')

add_h2(doc, 'Рядок 5 — Один чіткий CTA')
add_body(doc, 'Приклад: "If you\'re open to it, I\'d be happy to review {{CHANNELNAME}} and outline personalized terms based on your niche."')

doc.add_paragraph()
add_body(doc, 'Правило: не більше 150 слів. Немає списків у першому листі. Один CTA.')
doc.add_paragraph()

# 7. ПЛАН ДІЙ
add_h1(doc, '7. ПЛАН ДІЙ')
doc.add_paragraph()

add_table(doc,
    ['Пріоритет', 'Дія', 'Термін'],
    [
        ['Критично', 'Замінити "2B users" на "700M+ monthly users" у ВСІХ сіквенсах', 'До першої відправки'],
        ['Критично', 'Залишити один CTA на кожен лист — вибрати і зафіксувати', 'До першої відправки'],
        ['Критично', 'Вивести Launch 3 і Launch 5 з ротації', 'Зараз'],
        ['Важливо', 'Зробити 7 нішових версій Launch 10', 'Цього тижня'],
        ['Важливо', 'Запустити A/B тест: Launch 10 (Gain) vs Launch 4 (Loss)', 'Цього тижня'],
        ['Бажано', 'Розбити Launch 7 на 2 окремих сіквенси: MSN окремо, Safety Bundle окремо', 'Наступний тиждень'],
        ['Бажано', 'Стандартизувати цифри AIR (у деяких листах 20B переглядів, у деяких 30B — перевірити)', 'Наступний тиждень'],
    ]
)

doc.save('/home/user/Nataliia_Borska/MSN_Email_Audit.docx')
print('OK')
