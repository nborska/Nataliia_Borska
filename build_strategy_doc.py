from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Поля
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(16)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)

def add_body_mixed(doc, parts):
    # parts = list of (text, bold, italic)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic

def add_bullet(doc, text, bold_word=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)

def add_highlight(doc, text, color='blue'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if color == 'blue':
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif color == 'green':
        run.font.color.rgb = RGBColor(0x37, 0x86, 0x30)
    elif color == 'red':
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if hdr[i].paragraphs[0].runs:
            r = hdr[i].paragraphs[0].runs[0]
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        fill = 'F2F7FC' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            row[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            tc = row[c_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# ТИТУЛ
# ============================================================
add_title(doc, 'MSN DISTRIBUTION — СТРАТЕГІЧНИЙ РОЗБІР')
add_subtitle(doc, 'Корективи  ·  Пояснення термінів  ·  План дій')
add_subtitle(doc, 'На базі: Strategic Plan Leadership v1.0 + Product Guidelines AIR  |  Травень 2026')
add_divider(doc)

# ============================================================
# ЗМІСТ
# ============================================================
add_h1(doc, 'ЗМІСТ')
items = [
    '1. Корективи до плану',
    '2. Головний діагноз (проста версія)',
    '3. Пояснення термінів — усі 11 пунктів',
    '4. План по місяцях та тижнях',
    '5. Три головні експерименти',
    '6. Що зараз працює vs нові канали',
    '7. Що зупиняємо і чому',
    '8. Бюджет та стратегічне рішення',
]
for item in items:
    add_bullet(doc, item)
doc.add_paragraph()
add_divider(doc)

# ============================================================
# 1. КОРЕКТИВИ
# ============================================================
add_h1(doc, '1. КОРЕКТИВИ ДО ПЛАНУ')

add_h2(doc, 'AIR — "один з 4 офіційних партнерів"')
add_body(doc, 'У плані написано: AIR — один з 4 офіційних партнерів Microsoft.')
add_body(doc, 'Виправлення: AIR не "один з чотирьох" — партнерів більше, і AIR є одним із перших.')
add_body(doc, 'Як виправити у всіх матеріалах: замінити на "AIR — один з перших офіційних MSN Publishers" або "AIR — ранній офіційний партнер Microsoft". Це фактично точніше і звучить як конкурентна перевага (ранній вхід = стратегічна позиція), а не просто місце в черзі серед чотирьох.')

add_h2(doc, 'Кейси та контракти на лендингу')
add_body(doc, 'Лендинг air.io/ua/marketplace/msn-microsoft-network вже містить більше кейсів і контрактів ніж зазначено в плані.')
add_bullet(doc, 'Синхронізувати план з реальними цифрами з лендингу')
add_bullet(doc, 'Де план посилається на кейси — брати їх з лендингу, не занижувати кількість')
add_divider(doc)

# ============================================================
# 2. ДІАГНОЗ
# ============================================================
add_h1(doc, '2. ГОЛОВНИЙ ДІАГНОЗ (проста версія)')

add_highlight(doc, 'За 8 місяців прийшло 5 676 зацікавлених людей. Підписали угоду тільки 17.', 'blue')
add_body(doc, 'Проблема не в тому, що мало людей знають про MSN — проблема в тому, що 99.7% тих, хто вже зацікавився, "відвалились" до того, як стали клієнтами.')
add_body(doc, 'Тому план на червень — не "більше реклами", а "знайти де виходять люди і закрити дірки".')

doc.add_paragraph()
add_h3(doc, 'Три головні висновки з плану:')
add_table(doc,
    ['#', 'Висновок', 'Що з цим робити'],
    [
        ['1', 'Проблема у воронці, не в трафіку', 'Фікс SDR, шаблонів, лендингу'],
        ['2', 'Ринковий момент унікальний — 6-12 місяців', 'Позиція "Anti-Algorithm" поки вікно відкрите'],
        ['3', 'B2B-канал не використовується взагалі', 'LinkedIn до MCN — найвищий ROI'],
    ]
)
add_divider(doc)

# ============================================================
# 3. ТЕРМІНИ
# ============================================================
add_h1(doc, '3. ПОЯСНЕННЯ ТЕРМІНІВ')

# --- Фікс воронки ---
add_h2(doc, 'Фікс воронки')
add_body(doc, 'З плану: 3 168 "теплих" лідів охолонули — люди, що вже відреагували, але не стали клієнтами.')
add_body(doc, 'Головна причина: SDR відповідає через 1-2 дні замість менше 1 години.')
add_body(doc, 'Уяви: людина заповнила форму, чекає, настрій "так, хочу". Через 2 дні приходить відповідь — вона вже забула, охолола, зайнялась іншим. Гроші на рекламу витрачені, клієнта втрачено.')
doc.add_paragraph()
add_h3(doc, 'Воронка зараз:')
add_table(doc,
    ['Етап', 'Кількість', 'Конверсія'],
    [
        ['Входить (Leads)', '5 676', '—'],
        ['Теплі (MQL)', '3 168', '56%'],
        ['На дзвінок (SQL)', '190', '6% — ГОЛОВНА ДІРКА'],
        ['Підписали (Customers)', '17', '9%'],
    ]
)
add_body(doc, 'Фікс: відповідь протягом 1 години + сегментовані шаблони + оновлений лендинг.')

# --- Сегментовані шаблони ---
add_h2(doc, 'Сегментовані шаблони')
add_body(doc, 'Проблема зараз: один шаблон мейлу для всіх. Кулінарному і tech-каналу — один і той самий текст. Людина читає і думає "це не про мене".')
add_body(doc, 'Що треба: 7 різних шаблонів — по одному на кожну нішу:')
add_table(doc,
    ['Ніша', 'Кейс у шаблоні'],
    [
        ['Food & Cooking', 'Кулінарний канал: $410 → $1,300 за перший місяць'],
        ['DIY & Crafts', 'DIY канал: 17M переглядів за 3 місяці'],
        ['Tech & Science', 'Tech канал: 1.5M → 6.7M переглядів'],
        ['Health & Wellness', 'Lifehacks канал: +1.8M переглядів за 2 місяці'],
        ['Travel & Lifestyle', 'Кейс по подорожах (з лендингу)'],
        ['Business & Finance', 'Tech News: +$1,500 з back catalog'],
        ['Kids & Entertainment', 'Action Comedy: +677K щомісяця'],
    ]
)
add_body(doc, 'Принцип: кожен мейл показує кейс саме тієї ніші, до якої пишеш.')

# --- Лендинг ---
add_h2(doc, 'Що додати у лендинг (калькулятор — відкидаємо)')
add_body(doc, 'Що додати замість калькулятора:')
add_numbered(doc, 'Кейси по нішах — конкретні результати, розбиті за темою каналу (Food, DIY, Tech і т.д.)')
add_numbered(doc, 'FAQ — відповіді на страхи: "чи не конкурує з YouTube?", "чи залишу права?", "коли перший виплат?"')
add_numbered(doc, 'How it works — максимально 3 кроки: надіслав канал → ми перевірили → відео на MSN, гроші капають')
add_numbered(doc, 'Один чіткий CTA — не три кнопки, а одна: "Перевір, чи підходить твій канал (60 сек)"')
add_body(doc, 'Важливо з плану: кожна зайва кнопка знижує CTR головної на 10-20%. Менше варіантів = більше рішень.')

# --- RPM Report ---
add_h2(doc, 'RPM Report (серпень)')
add_body(doc, 'RPM = Revenue Per Mille = скільки грошей заробляється на 1000 переглядів.')
add_body(doc, 'З плану: в серпні виходить "State of MSN Monetization 2026" — знеособлений публічний звіт.')
add_body(doc, 'Що там буде:')
add_bullet(doc, 'Середній RPM по ринках і нішах')
add_bullet(doc, 'Динаміка Q3 2026 vs Q3 2025')
add_bullet(doc, 'Реальні діапазони доходу по типах каналів')
add_body(doc, 'Навіщо маркетингово: lead magnet для нових креаторів, соціальний доказ для тих хто сумнівається, контент для LinkedIn та Substack на весь серпень.')
add_body(doc, 'Чому серпень: рекламодавці повертаються після літа, RPM зростає в Q3-Q4 — саме час показати гарні цифри.')

# --- Discord ---
add_h2(doc, 'Discord присутність')
add_body(doc, 'З плану: перед підписанням лід шукає згадки AIR на Reddit, YouTube-коментах, Discord. Якщо знаходить 2-3 позитивні — підписує. Якщо знаходить 0 — не підписує.')
add_body(doc, 'Простими словами: людина майже готова підписати, але спочатку гуглить "AIR MSN відгуки". Якщо нічого — передумує. Якщо бачить 3 позитивні коментарі від живих людей — підписує.')
add_body(doc, 'Що робимо:')
add_bullet(doc, 'Заходимо у Discord-сервери де сидять YouTube-креатори')
add_bullet(doc, 'Відповідаємо на питання про монетизацію (80% корисність)')
add_bullet(doc, 'Природно згадуємо MSN тільки де це доречно (20%)')
add_body(doc, 'Ціль з плану: 5-10 публічних згадок AIR від реальних креаторів у різних точках інтернету.')

# --- TubeBuddy ---
add_h2(doc, 'Discord TubeBuddy — як нам там бути')
add_body(doc, 'TubeBuddy — популярний інструмент аналітики для YouTube. Їх Discord — один із найбільших хабів YouTube-креаторів.')
add_body(doc, 'Як бути там:')
add_numbered(doc, 'Реєструємо акаунт — нейтральне імя, не "AIR_official"')
add_numbered(doc, 'Беремо участь у розмовах про монетизацію та зростання каналу')
add_numbered(doc, 'Коли хтось скаржиться на падіння YouTube RPM (зараз таких багато) — це наш момент')
add_numbered(doc, 'Ділимось кейсами як корисною інформацією, не як рекламою')
add_body(doc, 'Правило: 1 органічний коментар на день цінніший за 10 рекламних.')
add_body(doc, 'Ціль: 5-10 розмов на тиждень → 2-3 заявки на місяць з цього каналу.')

# --- B2B ---
add_h2(doc, 'B2B масштабування')
add_body(doc, 'З плану — конкретні цифри:')
add_table(doc,
    ['Канал', 'Вартість клієнта', 'LTV:CAC'],
    [
        ['Cold email (зараз)', '$209', '2.9:1'],
        ['LinkedIn до MCN/агенцій (не запущений)', 'прогноз <$150', '20:1'],
        ['Affiliate MCN', 'прогноз ~$50', '60:1'],
    ]
)
add_body(doc, 'Простими словами: замість того щоб шукати кожного креатора окремо — шукаємо тих, хто вже має доступ до багатьох каналів.')
add_body(doc, 'Хто це:')
add_bullet(doc, 'Talent managers — ведуть 5-50 каналів')
add_bullet(doc, 'MCN-мережі — підписали 100+ каналів')
add_bullet(doc, 'YouTube-агентства — допомагають каналам рости')
add_body(doc, 'Чому ефективно: один договір з MCN = 10-50 каналів одразу, замість 50 окремих угод.')
add_highlight(doc, 'З плану: цей канал зараз взагалі не використовується — найбільший незайнятий важіль зростання.', 'red')

# --- Gain Frame ---
add_h2(doc, 'Gain Frame у cold outreach')
add_body(doc, 'З плану: рекомендується тестувати Loss Frame проти поточного підходу.')
add_table(doc,
    ['Тип фрейму', 'Приклад', 'Фокус'],
    [
        ['Gain Frame (зараз)', '"Твої відео можуть заробляти +30% на MSN"', 'Що ти ОТРИМАЄШ'],
        ['Loss Frame (тестуємо)', '"Твої 200 відео щомісяця залишають гроші на столі"', 'Що ти ЗАРАЗ ВТРАЧАЄШ'],
    ]
)
add_body(doc, 'Чому Loss Frame може бути ефективнішим: за дослідженням Канемана і Тверські — втрати психологічно важать у 2-2.5 рази сильніше за здобутки. Наш ICP прокидається не з думкою "хочу більше", а з думкою "AdSense знову впав".')
add_body(doc, 'Що робимо: A/B тест — половині пишемо Gain, половині Loss. Дивимось reply rate.')

# --- Бухгалтери ---
add_h2(doc, 'Бухгалтери — чому несподівано і як допоможуть')
add_body(doc, 'Чому несподівано: бухгалтер — не твій клієнт і не партнер. Але це людина, яка розмовляє з твоїми клієнтами щомісяця.')
add_body(doc, 'Логіка: хто першим знає, що YouTube-креатор хоче диверсифікувати дохід? Часто — його бухгалтер або фінансовий консультант.')
add_body(doc, 'Модель:')
add_bullet(doc, 'Знаходимо бухгалтерів/фін. консультантів що працюють з YouTube-креаторами, медіабізнесами')
add_bullet(doc, 'Пропонуємо реферальну програму: вони рекомендують AIR своїм клієнтам → отримують комісію')
add_bullet(doc, 'Для них легко: "Є нова можливість заробити ще 30% без нових відео — хочеш, познайомлю?"')
add_body(doc, 'Додатково: бухгалтери ходять на конференції Creator Economy і Digital-бізнесу. Вони — мультиплікатор довіри.')
add_body(doc, 'Що потрібно зробити: простий 1-pager для бухгалтерів + реферальні умови.')

# --- Substack ---
add_h2(doc, 'Substack Newsletter')
add_body(doc, 'Що таке Substack: платформа для email-розсилок. Можна безкоштовно збирати підписників по темах.')
add_body(doc, 'Чому Substack, а не просто email:')
add_bullet(doc, 'Є власна аудиторія — люди шукають авторів по темах і підписуються самі')
add_bullet(doc, 'Своя база підписників — ніхто не "відрізає" доступ (на відміну від Instagram)')
add_body(doc, 'Що публікуємо:')
add_bullet(doc, '1 лист на тиждень про монетизацію YouTube-креаторів')
add_bullet(doc, 'Кейси, тренди RPM, аналіз ринку')
add_bullet(doc, '80% корисний контент, 20% — наш оффер')
add_body(doc, 'Звязок з RPM report: "State of MSN Monetization 2026" — перший великий матеріал, який притягне підписників.')
add_body(doc, 'Метрика: 1,000 підписників за 3 місяці, open rate > 40%.')

# --- Flywheel ---
add_h2(doc, 'Органічний FLYWHEEL')
add_body(doc, 'Flywheel (маховик) — система де кожна дія посилює наступну і все починає крутитись само.')
add_body(doc, 'Звичайна логіка (лінійна): Реклама → Лід → Угода → Кінець')
add_body(doc, 'FLYWHEEL-логіка (циклічна):')
steps = [
    'Новий креатор підписується',
    'Отримує гарні результати',
    'Ми робимо кейс-стаді з ним',
    'Кейс іде в лендинг / email / LinkedIn / Substack',
    'Інший креатор бачить → приходить органічно',
    'Creator referral: $50-100 за приведеного',
    'Більше кейсів → flywheel розкручується далі',
]
for i, s in enumerate(steps):
    arrow = '↓' if i < len(steps) - 1 else ''
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(s)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    if arrow:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        p2.paragraph_format.space_after = Pt(1)
        p2.add_run(arrow)
doc.add_paragraph()
add_body(doc, 'З плану: Creator referral програма ($50-100 за referral) запускається в серпні — це механізм запуску flywheel.')
add_body(doc, 'Чому важливо запустити зараз: поки конкурентів мало на MSN — кожен кейс важить більше. Flywheel знижує вартість залучення клієнта з часом.')
add_divider(doc)

# ============================================================
# 4. ПЛАН ПО МІСЯЦЯХ
# ============================================================
add_h1(doc, '4. ПЛАН ПО МІСЯЦЯХ І ТИЖНЯХ')

# Червень
add_h2(doc, 'ЧЕРВЕНЬ — "Зупинити витік"')
add_body(doc, 'Суть: перш ніж витрачати більше на рекламу — залатати дірки у воронці.')
add_h3(doc, 'OKR (що вимірюємо):')
add_bullet(doc, 'MQL → SQL: з 6% до 8%')
add_bullet(doc, 'Reply rate cold email: з 10% до 15%+')
add_bullet(doc, 'Лендинг оновлений: кейси по нішах, FAQ, How-it-works, один CTA')
add_bullet(doc, 'Meta бюджет відновлено до $1,500/міс')
add_h3(doc, 'По тижнях:')
add_table(doc,
    ['Тиждень', 'Задача', 'Результат'],
    [
        ['Тиждень 1', 'Аудит воронки', 'Знаємо де конкретно "падають" ліди, є поточні метрики'],
        ['Тиждень 2', 'Фікс SDR', 'Правило відповіді за <1 годину, нові скрипти'],
        ['Тиждень 3', 'Нові шаблони мейлів', '7 сегментованих версій по нішах'],
        ['Тиждень 4', 'Оновлення лендингу', 'Кейси, FAQ, один CTA, прибрали зайве'],
    ]
)

# Липень
add_h2(doc, 'ЛИПЕНЬ — "Масштабувати що працює"')
add_body(doc, 'Суть: воронка підлатана → тепер натискаємо газ.')
add_h3(doc, 'OKR:')
add_bullet(doc, 'MQL → SQL: досягло 10%')
add_bullet(doc, '40+ SQL на місяць')
add_bullet(doc, 'LinkedIn outreach запущений: 15+ SQL/міс')
add_bullet(doc, 'Перша MCN-партнерська угода підписана')
add_h3(doc, 'По тижнях:')
add_table(doc,
    ['Тиждень', 'Задача', 'Результат'],
    [
        ['Тиждень 1', 'Loss Frame A/B тест', 'Запущений тест в cold email'],
        ['Тиждень 2', 'LinkedIn outreach', 'Старт контактів з MCN і talent managers'],
        ['Тиждень 3', 'Масштабування Meta', 'Той креатив що конвертує, A/B тільки 1 елемент'],
        ['Тиждень 4', 'Discord присутність', 'Перші 5 органічних згадок AIR'],
    ]
)

# Серпень
add_h2(doc, 'СЕРПЕНЬ — "Запустити Referral Loop"')
add_body(doc, 'Суть: compound growth — канали що ростуть без постійного бюджету.')
add_h3(doc, 'OKR:')
add_bullet(doc, 'Дохід $1,500+/міс (на шляху до $1,700)')
add_bullet(doc, '5-6 нових клієнтів на місяць')
add_bullet(doc, 'Creator referral програма запущена')
add_bullet(doc, '"State of MSN Monetization 2026" опублікований')
add_h3(doc, 'По тижнях:')
add_table(doc,
    ['Тиждень', 'Задача', 'Результат'],
    [
        ['Тиждень 1', 'Підготовка RPM-звіту', 'Зібрані дані від партнерів'],
        ['Тиждень 2', 'Referral + Substack', 'Запуск $50-100 за приведеного, перший лист'],
        ['Тиждень 3', 'Публікація звіту', '"State of MSN Monetization 2026" виходить'],
        ['Тиждень 4', 'Аналіз кварталу', 'Що продовжуємо, що зупиняємо, план Q4'],
    ]
)
add_divider(doc)

# ============================================================
# 5. ЕКСПЕРИМЕНТИ
# ============================================================
add_h1(doc, '5. ТРИ ГОЛОВНІ ЕКСПЕРИМЕНТИ')
add_body(doc, 'ICE Score = Impact x Confidence x Ease / 3. Чим вище — тим вищий пріоритет. Шкала 1-10.')
doc.add_paragraph()

add_h2(doc, 'Експеримент 1 — SDR відповідь за <1 годину')
add_highlight(doc, 'ICE Score: 8.7 — найвищий пріоритет', 'blue')
add_table(doc,
    ['Параметр', 'Деталь'],
    [
        ['Проблема', 'SDR відповідає через 1-2 дні → лід охолов → угода не закрилась'],
        ['Що тестуємо', 'Відповідаємо протягом 1 години після заявки'],
        ['Механіка', 'Половина лідів отримує відповідь за <1 год, половина — як зараз'],
        ['Метрика успіху', 'Зростання конверсії MQL → SQL у групі з <1 год відповіддю'],
        ['Ризик', 'Потрібен SDR-ресурс, який реально моніторить заявки'],
    ]
)

add_h2(doc, 'Експеримент 2 — Loss Frame A/B тест')
add_highlight(doc, 'ICE Score: 8.3 — другий за пріоритетом', 'blue')
add_table(doc,
    ['Варіант', 'Текст', 'Фокус'],
    [
        ['A — Gain (зараз)', '"Твої відео можуть заробляти +30% на MSN"', 'Що отримаєш'],
        ['B — Loss (тест)', '"Твої 200 відео щомісяця залишають гроші на столі"', 'Що втрачаєш'],
    ]
)
add_body(doc, 'Метрика: reply rate і конверсія по кожній групі.')
add_body(doc, 'Якщо Loss Frame виграє: переписуємо всі мейли, рекламу і частину лендингу.')

add_h2(doc, 'Експеримент 3 — Revenue Calculator на лендингу')
add_highlight(doc, 'ICE Score: 7.7 — ВІДКЛАДЕНО за рішенням команди', 'red')
add_body(doc, 'Замість калькулятора — конкретні приклади доходу по нішах з реальних кейсів.')
add_divider(doc)

# ============================================================
# 6. КАНАЛИ
# ============================================================
add_h1(doc, '6. ЩО ЗАРАЗ ПРАЦЮЄ VS НОВІ КАНАЛИ')

add_h2(doc, 'Працює → масштабуємо:')
add_table(doc,
    ['Канал', 'Поточна метрика', 'Що змінюємо'],
    [
        ['Cold email outreach', '$209/клієнт, LTV:CAC 2.9:1', 'Сегментуємо по 7 нішах, Loss Frame A/B, SDR < 1 год'],
        ['Meta реклама', 'Є конверсії, останній креатив — найкращий', 'Відновлюємо до $1,500/міс, масштабуємо той що конвертує, тестуємо лише 1 елемент за раз'],
    ]
)

add_h2(doc, 'Нові канали — що там робимо:')
add_table(doc,
    ['Канал', 'Коли', 'Що конкретно'],
    [
        ['LinkedIn → MCN/agencies', 'Липень', 'Прямий outreach, 1 договір = 10-50 каналів, LTV:CAC прогноз 20:1'],
        ['Discord', 'Червень-Липень', 'Organic presence, peer validation, 5-10 згадок AIR'],
        ['Substack newsletter', 'Серпень', 'Weekly newsletter + RPM-звіт як перший великий матеріал'],
        ['Creator Referral', 'Серпень', '$50-100 за кожного приведеного — запускає flywheel'],
        ['Бухгалтери/фін. консультанти', 'Q4 (після flywheel)', 'Реферальна програма, 1-pager, вхід через конференції'],
    ]
)
add_divider(doc)

# ============================================================
# 7. ЗУПИНЯЄМО
# ============================================================
add_h1(doc, '7. ЩО ЗУПИНЯЄМО І ЧОМУ')
add_table(doc,
    ['Канал', 'Чому зупиняємо', 'Цифра'],
    [
        ['Google Search Ads', 'Надто дорого', '$2,598 за клієнта — у 12 разів дорожче ніж cold email'],
        ['TikTok Ads', 'Підтверджений провал', '$143 CPL — не окупається'],
    ]
)
add_divider(doc)

# ============================================================
# 8. БЮДЖЕТ І СТРАТЕГІЧНЕ РІШЕННЯ
# ============================================================
add_h1(doc, '8. БЮДЖЕТ ТА СТРАТЕГІЧНЕ РІШЕННЯ')

add_h2(doc, 'Бюджет на 3 місяці')
add_highlight(doc, 'Загальний запит: $9,380  |  Прогноз ROMI: 5.4:1', 'blue')
add_table(doc,
    ['Стаття', 'Сума'],
    [
        ['Meta Ads', '$5,000'],
        ['LinkedIn Ads / outreach інструменти', '$1,400'],
        ['Інше (інструменти, контент, referral)', '$3,000'],
        ['РАЗОМ', '$9,380'],
    ]
)

add_h2(doc, 'Важливе стратегічне рішення (на обговорення)')
add_table(doc,
    ['Параметр', 'Деталь'],
    [
        ['Зараз', '"Ми допоможемо розповсюдити твої відео на MSN"'],
        ['Пропозиція', '"Ми захищаємо тебе від залежності від YouTube-алгоритму" (Anti-Algorithm)'],
        ['Чому резонує', 'YouTube Algorithm Apocalypse 2025-2026: падіння переглядів 25-90%, 60% creators шукають альтернативи'],
        ['Ризик', 'Змінює весь messaging, потребує rewrite лендингу, мейлів, реклами'],
        ['Рекомендація', 'Обговорити з керівництвом до старту червня'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Документ підготовлений на базі MSN Strategic Plan Leadership v1.0 та AIR Product Guidelines')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

doc.save('/home/user/Nataliia_Borska/MSN_Strategy_Breakdown.docx')
print('OK')
